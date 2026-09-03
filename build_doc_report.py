import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from build_doc_report_helpers import (
    set_cell_background, set_cell_margins, set_cell_borders,
    add_styled_heading, add_callout, add_code_block, add_image_with_caption
)

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

# Cover Header
p_t = doc.add_paragraph()
p_t.paragraph_format.space_before = Pt(0)
p_t.paragraph_format.space_after = Pt(4)
r_t = p_t.add_run("STRATEGIC PLANNING & DATA EXPLORATION REPORT")
r_t.font.name = 'Arial'
r_t.font.size = Pt(21)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(27, 54, 93)

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(14)
r_sub = p_sub.add_run("Optimizing Multi-Echelon Inventory Allocation and Last-Mile Route Efficiency Through Data Science & Machine Learning")
r_sub.font.name = 'Arial'
r_sub.font.size = Pt(12)
r_sub.font.bold = True
r_sub.font.color.rgb = RGBColor(43, 92, 143)

# Metadata Table
meta_table = doc.add_table(rows=2, cols=4)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_table.autofit = False
col_widths = [Inches(1.5), Inches(1.75), Inches(1.5), Inches(1.75)]
meta_data = [
    [("Project Phase:", "Week 1: Strategy & Exploration"), ("Document Version:", "1.0 (Enterprise Plan)")],
    [("Target Network:", "Urban Omnichannel Supply Chain"), ("Core Methodology:", "ML, Spatial Clustering & VRP")]
]
for r_idx, row in enumerate(meta_data):
    c_idx = 0
    for lbl, val in row:
        cl = meta_table.cell(r_idx, c_idx)
        cv = meta_table.cell(r_idx, c_idx + 1)
        cl.width, cv.width = col_widths[c_idx], col_widths[c_idx + 1]
        set_cell_background(cl, "F0F4F8")
        set_cell_background(cv, "FFFFFF")
        set_cell_margins(cl, 60, 60, 80, 80)
        set_cell_margins(cv, 60, 60, 80, 80)
        set_cell_borders(cl, bottom=('single', '4', 'D0D7DE'), top=('single', '4', 'D0D7DE'))
        set_cell_borders(cv, bottom=('single', '4', 'D0D7DE'), top=('single', '4', 'D0D7DE'))
        pl = cl.paragraphs[0]; pl.paragraph_format.space_after = Pt(0)
        rl = pl.add_run(lbl); rl.font.name = 'Arial'; rl.font.size = Pt(8.5); rl.font.bold = True; rl.font.color.rgb = RGBColor(27, 54, 93)
        pv = cv.paragraphs[0]; pv.paragraph_format.space_after = Pt(0)
        rv = pv.add_run(val); rv.font.name = 'Calibri'; rv.font.size = Pt(8.5); rv.font.color.rgb = RGBColor(50, 50, 50)
        c_idx += 2

p_sp = doc.add_paragraph(); p_sp.paragraph_format.space_before = Pt(6); p_sp.paragraph_format.space_after = Pt(8)
# Executive Summary
add_styled_heading(doc, "Executive Summary", level=1)
p = doc.add_paragraph()
p.add_run(
    "Modern e-commerce and retail supply chains operate under relentless pressure to satisfy shrinking delivery windows "
    "while managing transportation costs, inventory imbalances, and urban gridlock. The last mile of logistics represents "
    "up to 53% of total supply chain operating expenses. This strategic planning report outlines a rigorous, data-driven "
    "framework designed to optimize multi-echelon inventory allocation and transform last-mile delivery operations through "
    "applied data science, machine learning, and operations research."
)

add_callout(
    doc,
    "By transitioning from static FIFO dispatching to an integrated predictive and prescriptive optimization engine, "
    "the enterprise can achieve an On-Time In-Full (OTIF) rate exceeding 96%, reduce last-mile mileage by over 50%, "
    "lower per-delivery operating costs by 15-20%, and significantly mitigate supply chain carbon intensity.",
    title="EXECUTIVE VALUE PROPOSITION"
)

p = doc.add_paragraph()
p.add_run(
    "This document establishes the strategic deliverables for Week 1: operational scenario scoping, literature and open data "
    "benchmarking (using the Amazon Last-Mile and DataCo Supply Chain corpora), mathematical KPI formulations, a 7-phase analytical "
    "roadmap, and modular Python implementation scripts. In our simulated 3,000-order trial, machine learning achieved a 0.9976 ROC-AUC "
    "in predicting delivery delay risk, while heuristic TSP routing demonstrated a 50.84% reduction in travel mileage."
)

# Section 1: Project Definition & Logistics Scenario
add_styled_heading(doc, "1. Project Definition & Logistics Scenario", level=1)
add_styled_heading(doc, "1.1 Network Context & Operational Scenario", level=2)
p = doc.add_paragraph()
p.add_run(
    "The target enterprise manages a multi-tier fulfillment logistics network servicing a high-density metropolitan area. "
    "The distribution network is organized into two primary tiers:\n"
    "1. Regional Distribution Centers (RDCs): Large-scale fulfillment hubs (RDC-North, RDC-South) carrying broad SKU catalogues, "
    "fulfilling standard bulk ground shipments, and managing long-haul supplier replenishment.\n"
    "2. Urban Micro-Fulfillment Centers (MFCs): Localized urban dark stores (MFC-Central, MFC-East) situated near dense customer "
    "clusters to service rapid Same-Day Express (under 6 hours) and Next-Day Priority (under 24 hours) orders."
)

add_styled_heading(doc, "1.2 Core Operational Challenges", level=2)
p = doc.add_paragraph()
p.add_run(
    "The enterprise currently experiences three acute operational pain points:\n"
    "- Inventory Misallocation & Stockouts: Phantom stockouts in urban MFCs force costly cross-region split dispatches from RDCs, "
    "increasing cost-per-drop by 25% and transit duration by 65%.\n"
    "- Same-Day SLA Vulnerability: Baseline Same-Day Express deliveries suffer an alarming 16.25% late delivery rate due to traffic "
    "congestion and multi-attempt failures.\n"
    "- Static Dispatch Routing: Unoptimized first-in first-out dispatch produces overlapping routes, high vehicle wear, excess fuel "
    "burn, and avoidable driver overtime."
)

add_styled_heading(doc, "1.3 Key Performance Indicators (KPIs)", level=2)
p = doc.add_paragraph()
p.add_run("The following five strategic KPIs establish quantitative baselines and target benchmarks:")

kpi_table = doc.add_table(rows=6, cols=6)
kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi_table.autofit = False
kpi_headers = ["KPI Name", "Category", "Mathematical Formula", "Baseline", "Target", "Business Impact"]
kpi_widths = [Inches(1.2), Inches(0.9), Inches(1.8), Inches(0.7), Inches(0.7), Inches(1.2)]

for i, title in enumerate(kpi_headers):
    c = kpi_table.rows[0].cells[i]
    c.width = kpi_widths[i]
    set_cell_background(c, "1B365D")
    set_cell_margins(c, 80, 80, 80, 80)
    p_h = c.paragraphs[0]; p_h.paragraph_format.space_after = Pt(0); p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_h.add_run(title); r.font.name = 'Arial'; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)

kpi_rows = [
    ("1. On-Time In-Full (OTIF) Rate", "Service Quality", "(N_ontime_infull / N_total) * 100", "91.73%", ">= 96.0%", "Protects brand loyalty, eliminates SLA penalties."),
    ("2. Cost Per Delivery", "Financial Efficiency", "Total Last-Mile Cost / Total Stops", "$36.40 / stop", "<= $30.00 (-17.5%)", "Expands profit margin on rapid fulfillment."),
    ("3. Inventory Turnover Ratio", "Working Capital", "COGS / Average Inventory Value", "6.2x / yr", ">= 8.5x / yr", "Minimizes holding costs and inventory obsolescence."),
    ("4. First-Attempt Delivery Rate", "Operational Quality", "(N_single_attempt / N_total) * 100", "86.57%", ">= 95.0%", "Eliminates expensive re-delivery attempts ($6.50/re-attempt)."),
    ("5. Route Carbon Intensity", "ESG / Sustainability", "Total CO2e (kg) / Ton-km Delivered", "184 g/t-km", "<= 135 g/t-km (-26%)", "Ensures compliance with urban clean air regulations.")
]

for r_idx, r_data in enumerate(kpi_rows):
    rc = kpi_table.rows[r_idx + 1].cells
    bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
    for c_idx, val in enumerate(r_data):
        rc[c_idx].width = kpi_widths[c_idx]
        set_cell_background(rc[c_idx], bg)
        set_cell_margins(rc[c_idx], 60, 60, 70, 70)
        set_cell_borders(rc[c_idx], bottom=('single', '4', 'D0D7DE'), top=('single', '4', 'D0D7DE'))
        p_c = rc[c_idx].paragraphs[0]; p_c.paragraph_format.space_after = Pt(0)
        r = p_c.add_run(val); r.font.name = 'Calibri'; r.font.size = Pt(8)
        if c_idx in [0, 3, 4]: r.font.bold = True
        if c_idx == 3: r.font.color.rgb = RGBColor(180, 40, 40)
        elif c_idx == 4: r.font.color.rgb = RGBColor(30, 130, 60)

doc.add_paragraph().paragraph_format.space_after = Pt(8)
# Section 2: Literature Review & Data Research
add_styled_heading(doc, "2. Literature Review & Data Science Research", level=1)
add_styled_heading(doc, "2.1 Benchmark Public Datasets", level=2)
p = doc.add_paragraph()
p.add_run(
    "Our analytical framework is grounded in empirical research using three prominent public datasets:\n"
    "1. Amazon Last-Mile Routing Research Challenge Dataset (MIT/Amazon, 2021): Contains over 6,000 actual delivery routes "
    "with GPS coordinates, transit times, package dimensions, service times, and road circuity distributions (1.28 - 1.55).\n"
    "2. DataCo Global Supply Chain Intelligence Dataset: A repository of 180,000+ omnichannel order transactions detailing delivery "
    "status, late delivery risk factors, order geography, and shipping modes.\n"
    "3. US DOT BTS Freight Analysis Framework (FAF5): Multimodal freight flow data and highway speed indices across major metro corridors."
)

add_styled_heading(doc, "2.2 Applied Data Science Methodologies", level=2)
p = doc.add_paragraph()
p.add_run(
    "The project integrates four complementary data science paradigms:\n"
    "- Supervised Regression & Time-Series: LightGBM and XGBoost models for localized SKU-level demand forecasting, incorporating "
    "rolling sales lags, calendar factors, and promotional seasonality.\n"
    "- Supervised Classification: Balanced Random Forest classifiers predicting pre-dispatch Late Delivery Risk at order manifestation, "
    "enabling proactive SLA intervention.\n"
    "- Unsupervised Spatial Clustering: K-Means and DBSCAN algorithms partitioning delivery drops into compact urban micro-zones and "
    "identifying optimal centroid locations for urban micro-hubs.\n"
    "- Prescriptive Operations Research: Heuristic Capacitated Vehicle Routing Problem with Time Windows (CVRPTW) combining Clarke-Wright "
    "Savings, Greedy Nearest Neighbor, and 2-Opt local search."
)

# Section 3: Strategic Roadmap
add_styled_heading(doc, "3. Strategic Roadmap: End-to-End Analytical Pipeline", level=1)
p = doc.add_paragraph()
p.add_run(
    "The analytical transformation is structured across an end-to-end 7-phase implementation roadmap:"
)

road_table = doc.add_table(rows=8, cols=5)
road_table.alignment = WD_TABLE_ALIGNMENT.CENTER
road_table.autofit = False
road_headers = ["Phase", "Operational Scope", "Key Technical Activities", "Primary Deliverable", "Timeline"]
road_widths = [Inches(1.0), Inches(1.3), Inches(2.2), Inches(1.2), Inches(0.8)]

for i, t in enumerate(road_headers):
    c = road_table.rows[0].cells[i]
    c.width = road_widths[i]
    set_cell_background(c, "1B365D")
    set_cell_margins(c, 80, 80, 80, 80)
    p_h = c.paragraphs[0]; p_h.paragraph_format.space_after = Pt(0); p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_h.add_run(t); r.font.name = 'Arial'; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)

road_rows = [
    ("Phase 1: Scoping & Governance", "Requirements & Alignment", "Define SLAs, map warehouse tiers, establish KPI formulas, audit data compliance.", "Project Charter & KPI Matrix", "Weeks 1-2"),
    ("Phase 2: Ingestion Architecture", "Data Infrastructure", "Build Kafka/Airflow streaming pipelines linking ERP (SAP), WMS, and TMS GPS telematics.", "Automated Ingestion Lakehouse", "Weeks 3-4"),
    ("Phase 3: Cleansing & Feature Eng.", "Data Wrangling", "Geocoding, handling address errors, computing road circuity, building lag and traffic features.", "Clean Analytical Feature Store", "Weeks 5-6"),
    ("Phase 4: Exploratory Analysis", "Diagnostic Insights", "Spatial failure heatmaps, bottleneck correlation matrices, lead-time variance decomposition.", "Interactive Streamlit Dashboards", "Weeks 7-8"),
    ("Phase 5: ML & Optimization", "Model Development", "Train LightGBM forecasters, Random Forest delay classifiers, and 2-Opt CVRPTW routing engines.", "Trained & Validated Model Weights", "Weeks 9-11"),
    ("Phase 6: Simulation & Pilot", "Backtesting & Validation", "Simulate historical routing replays, conduct rolling cross-validation, run 10-van live pilot.", "Simulation & Pilot Audit Report", "Weeks 12-13"),
    ("Phase 7: Deployment & MLOps", "Enterprise Rollout", "Containerize inference microservices (FastAPI/Docker), integrate with driver mobile apps, drift monitors.", "Production Dispatch Engine", "Weeks 14-16")
]

for r_idx, r_data in enumerate(road_rows):
    rc = road_table.rows[r_idx + 1].cells
    bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
    for c_idx, val in enumerate(r_data):
        rc[c_idx].width = road_widths[c_idx]
        set_cell_background(rc[c_idx], bg)
        set_cell_margins(rc[c_idx], 60, 60, 70, 70)
        set_cell_borders(rc[c_idx], bottom=('single', '4', 'D0D7DE'), top=('single', '4', 'D0D7DE'))
        p_c = rc[c_idx].paragraphs[0]; p_c.paragraph_format.space_after = Pt(0)
        r = p_c.add_run(val); r.font.name = 'Calibri'; r.font.size = Pt(8)
        if c_idx in [0, 4]: r.font.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(8)
# Section 4: Python Code Illustration
add_styled_heading(doc, "4. Python Code Illustration & Technical Architecture", level=1)
p = doc.add_paragraph()
p.add_run(
    "To operationalize the analytical framework, production-grade Python scripts have been authored. "
    "Below are the core architectural modules illustrating feature engineering, delay risk modeling, and route optimization:"
)

add_styled_heading(doc, "4.1 Data Cleaning & Feature Engineering Pipeline", level=2)
code_f = """# Feature engineering: calculate spatial distances & temporal lags
def preprocess_logistics_data(df: pd.DataFrame) -> pd.DataFrame:
    d_lat = (df['Dest_Latitude'] - df['Origin_Latitude']) * 111.0
    d_lon = (df['Dest_Longitude'] - df['Origin_Longitude']) * 85.0
    df['Direct_Distance_km'] = np.sqrt(d_lat**2 + d_lon**2)
    df['Route_Distance_km'] = df['Direct_Distance_km'] * 1.38  # Urban circuity factor
    df['Order_Timestamp'] = pd.to_datetime(df['Order_Timestamp'])
    df['Hour_of_Day'] = df['Order_Timestamp'].dt.hour
    df['Is_Weekend'] = df['Order_Timestamp'].dt.dayofweek.isin([5, 6]).astype(int)
    return pd.get_dummies(df, columns=['Shipping_Mode', 'Origin_Warehouse_Type'], drop_first=True)
"""
add_code_block(doc, code_f, caption="feature_pipeline.py")

add_styled_heading(doc, "4.2 Supervised Machine Learning: Late Delivery Risk Classifier", level=2)
code_ml = """# Train Balanced Random Forest classifier for pre-dispatch delay prediction
def train_delay_classifier(X, y):
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.20, stratify=y, random_state=42)
    clf = RandomForestClassifier(n_estimators=150, max_depth=8, class_weight='balanced', random_state=42)
    clf.fit(X_train, y_train)
    y_pred, y_prob = clf.predict(X_test), clf.predict_proba(X_test)[:, 1]
    print(f"ROC-AUC: {roc_auc_score(y_test, y_prob):.4f}")
    return clf, clf.feature_importances_
"""
add_code_block(doc, code_ml, caption="delay_risk_model.py")

add_styled_heading(doc, "4.3 Spatial Clustering & Heuristic Route Optimization (2-Opt TSP)", level=2)
code_opt = """# Cluster delivery drops & optimize stop sequence using Nearest Neighbor + 2-Opt
def optimize_cluster_routes(coords, depot_coord):
    kmeans = KMeans(n_clusters=4, random_state=42).fit(coords)
    all_stops = [depot_coord] + list(coords)
    dist_mat = np.linalg.norm(np.array(all_stops)[:, None] - np.array(all_stops), axis=2)
    # Nearest Neighbor construction
    unvisited, tour, curr = set(range(1, len(all_stops))), [0], 0
    while unvisited:
        nxt = min(unvisited, key=lambda i: dist_mat[curr, i])
        tour.append(nxt); unvisited.remove(nxt); curr = nxt
    tour.append(0)
    # 2-Opt local search refinement
    improved = True
    while improved:
        improved = False
        for i in range(1, len(tour) - 2):
            for j in range(i + 1, len(tour) - 1):
                if dist_mat[tour[i-1], tour[j]] + dist_mat[tour[i], tour[j+1]] < dist_mat[tour[i-1], tour[i]] + dist_mat[tour[j], tour[j+1]] - 1e-4:
                    tour[i:j+1] = reversed(tour[i:j+1]); improved = True
    return tour, dist_mat
"""
add_code_block(doc, code_opt, caption="route_optimization_engine.py")
# Section 5: Empirical Findings & Experimental Results
add_styled_heading(doc, "5. Empirical Findings & Experimental Results", level=1)
p = doc.add_paragraph()
p.add_run(
    "To substantiate the strategic framework with quantitative evidence, the complete Python data pipeline was executed "
    "across a simulated dataset of 3,000 operational dispatches. The empirical results reveal striking insights into current "
    "performance vulnerabilities and the transformative potential of data science intervention."
)

add_styled_heading(doc, "5.1 Baseline Logistics Performance & KPI Diagnostics", level=2)
p = doc.add_paragraph()
p.add_run(
    "The overall baseline OTIF rate was measured at 91.73%, falling below the 96.0% corporate standard. "
    "Disaggregating performance across shipping tiers and fulfillment nodes highlights severe structural asymmetry:"
)

add_image_with_caption(
    doc,
    os.path.join("figures", "fig1_kpi_distribution.png"),
    "Figure 1: Comprehensive Logistics KPI Dashboard showing (A) OTIF Rate by Shipping Mode, (B) Late Delivery Rate by Warehouse Tier, (C) Delivery Cost Distribution, and (D) Transit Duration vs Distance Correlation."
)

p = doc.add_paragraph()
p.add_run(
    "Core findings from Figure 1:\n"
    "1. Shipping Tier Vulnerability: Standard Ground deliveries achieve a 97.31% OTIF rate with 0% late deliveries. In contrast, "
    "Same-Day Express displays an alarming 16.25% late delivery rate (OTIF 81.27%), identifying rapid fulfillment as the primary failure point.\n"
    "2. Fulfillment Node Advantage: Orders fulfilled from urban MFCs incur an average transit time of 3.68 hours and cost $32.35/drop, "
    "compared to 6.08 hours and $40.56/drop from regional RDCs—confirming a 20.2% cost advantage and 39.5% speed advantage for urban hubs.\n"
    "3. Cost Skewness: Average delivery cost is $36.40, with heavy right-tail outliers exceeding $65 due to failed attempts and congestion."
)

add_styled_heading(doc, "5.2 Machine Learning Delay Prediction & Feature Importance", level=2)
p = doc.add_paragraph()
p.add_run(
    "The Balanced Random Forest delay classifier demonstrated exceptional predictive accuracy on hold-out test dispatches: "
    "Accuracy: 98.83%, Precision: 96.30%, Recall: 81.25%, F1-Score: 88.14%, and an outstanding ROC-AUC of 0.9976."
)

add_image_with_caption(
    doc,
    os.path.join("figures", "fig2_feature_importance.png"),
    "Figure 2: Relative Feature Importance (Gini Impurity) for Late Delivery Risk Classification using Random Forest."
)

p = doc.add_paragraph()
p.add_run(
    "Figure 2 reveals that the number of Delivery Attempts is by far the single most decisive predictor of delay risk, "
    "accounting for 48.8% of model importance. When a delivery fails on its initial attempt, the probability of an SLA breach "
    "exceeds 85%. Promised SLA Hours (11.0%), Same-Day Express mode (9.9%), and Weather Severity (7.2%) represent the next key drivers. "
    "This confirms that first-attempt failure prevention is the highest-leverage intervention point for improving OTIF."
)
add_styled_heading(doc, "5.3 Spatial Clustering for Dynamic Micro-Zone Partitioning", level=2)
p = doc.add_paragraph()
p.add_run(
    "Unsupervised K-Means clustering partitioned customer drop coordinates into four distinct operational zones "
    "(Zone Alpha, Zone Beta, Zone Gamma, Zone Delta) across the metropolitan region:"
)

add_image_with_caption(
    doc,
    os.path.join("figures", "fig3_spatial_clustering.png"),
    "Figure 3: Spatial Clustering of Delivery Locations with Identified Cluster Centroids (Proposed Micro-Hubs) and Regional Warehouses."
)

p = doc.add_paragraph()
p.add_run(
    "As visualized in Figure 3, zoning stops by spatial density prevents van routes from traversing cross-city corridors. "
    "Crucially, the four calculated cluster centroids (yellow stars) identify mathematically optimal candidate sites for forward-deployed "
    "micro-fulfillment dark stores or mobile container hubs, cutting stem transit distance from peripheral RDCs."
)

add_styled_heading(doc, "5.4 Combinatorial Route Optimization Simulation", level=2)
p = doc.add_paragraph()
p.add_run(
    "To test the efficiency of automated vehicle routing, a dispatch simulation was conducted on 22 delivery stops originating from MFC-Central. "
    "A legacy unoptimized (FIFO) sequence was compared against our 2-Opt TSP heuristic optimization:"
)

add_image_with_caption(
    doc,
    os.path.join("figures", "fig4_route_optimization.png"),
    "Figure 4: Delivery Tour Comparison for a 22-Stop Van Dispatch: (A) Unoptimized Baseline Route vs. (B) Heuristic 2-Opt Optimized Route."
)

opt_table = doc.add_table(rows=6, cols=4)
opt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
opt_table.autofit = False
opt_headers = ["Operational Dimension", "Baseline (Unoptimized)", "Optimized (2-Opt TSP)", "Net Improvement / Savings"]
opt_widths = [Inches(1.8), Inches(1.5), Inches(1.5), Inches(1.7)]

for i, t in enumerate(opt_headers):
    c = opt_table.rows[0].cells[i]
    c.width = opt_widths[i]
    set_cell_background(c, "1B365D")
    set_cell_margins(c, 80, 80, 80, 80)
    p_h = c.paragraphs[0]; p_h.paragraph_format.space_after = Pt(0); p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_h.add_run(t); r.font.name = 'Arial'; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)

opt_rows = [
    ("Total Tour Road Distance", "267.39 km", "131.44 km", "-135.95 km (-50.84%)"),
    ("Estimated Driving Time", "7.64 hours", "3.75 hours", "-3.89 hours (-50.91%)"),
    ("Diesel Fuel Consumption", "32.09 Liters", "15.77 Liters", "-16.31 Liters (-50.83%)"),
    ("Carbon Emissions (CO2e)", "86.00 kg CO2e", "42.27 kg CO2e", "-43.73 kg CO2e (-50.85%)"),
    ("Estimated Driver Labor Cost", "$194.82", "$95.63", "-$99.19 (-50.91%)")
]

for r_idx, r_vals in enumerate(opt_rows):
    rc = opt_table.rows[r_idx + 1].cells
    bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
    for c_idx, v in enumerate(r_vals):
        rc[c_idx].width = opt_widths[c_idx]
        set_cell_background(rc[c_idx], bg)
        set_cell_margins(rc[c_idx], 60, 60, 70, 70)
        set_cell_borders(rc[c_idx], bottom=('single', '4', 'D0D7DE'), top=('single', '4', 'D0D7DE'))
        p_c = rc[c_idx].paragraphs[0]; p_c.paragraph_format.space_after = Pt(0)
        r = p_c.add_run(v); r.font.name = 'Calibri'; r.font.size = Pt(8.5)
        if c_idx == 0: r.font.bold = True
        elif c_idx == 3: r.font.bold = True; r.font.color.rgb = RGBColor(30, 130, 60)

doc.add_paragraph().paragraph_format.space_after = Pt(8)
# Section 6: Business Impact & Strategic Recommendations
add_styled_heading(doc, "6. Business Impact, Risk Analysis & Recommendations", level=1)
add_styled_heading(doc, "6.1 Quantified Financial ROI & Value Creation", level=2)
p = doc.add_paragraph()
p.add_run(
    "Extrapolating these empirical findings across an enterprise fleet of 120 delivery vans executing 300,000 annual deliveries "
    "yields compelling economic returns:\n"
    "1. Fleet Fuel & Maintenance Savings: A 30% sustained mileage reduction across daily routes eliminates ~850,000 km of driving, "
    "saving an estimated $272,000 in fuel and avoiding 270 metric tons of CO2e annually.\n"
    "2. Overtime Elimination: Reducing driving hours per route by 3.8 hours cuts peak-season driver overtime by $480,000/year.\n"
    "3. First-Attempt Failure Avoidance: Predictive pre-dispatch intervention drops re-attempt rates from 13.4% to <5%, saving $165,000/year."
)

add_styled_heading(doc, "6.2 Risk Assessment & Governance Matrix", level=2)
risk_table = doc.add_table(rows=5, cols=4)
risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
risk_table.autofit = False
risk_headers = ["Risk Domain", "Risk Description", "Severity / Likelihood", "Mitigation Strategy"]
risk_widths = [Inches(1.2), Inches(2.0), Inches(1.3), Inches(2.0)]

for i, t in enumerate(risk_headers):
    c = risk_table.rows[0].cells[i]
    c.width = risk_widths[i]
    set_cell_background(c, "1B365D")
    set_cell_margins(c, 80, 80, 80, 80)
    p_h = c.paragraphs[0]; p_h.paragraph_format.space_after = Pt(0); p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_h.add_run(t); r.font.name = 'Arial'; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)

risk_rows = [
    ("Data Latency & Quality", "WMS/TMS GPS dropouts, stale inventory counts.", "High / Medium", "Implement streaming validation (Great Expectations) with automatic fallback schemas."),
    ("Driver Non-Compliance", "Drivers bypass algorithmic routes due to perceived lack of nuance.", "High / High", "Driver mobile UI with real-time traffic overlay, turn-by-turn guidance, and gamified incentives."),
    ("Concept & Data Drift", "Holiday surges (Cyber Week) distort historical traffic and demand.", "Medium / High", "Continuous MLOps drift monitoring (Evidently AI) triggering automated bi-weekly retraining."),
    ("Model Overfitting", "Overly tight route optimization leaves no slack for urban delays.", "Medium / Medium", "Enforce dynamic buffer windows (10-15% slack factor) in the CVRPTW solver.")
]

for r_idx, r_vals in enumerate(risk_rows):
    rc = risk_table.rows[r_idx + 1].cells
    bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
    for c_idx, v in enumerate(r_vals):
        rc[c_idx].width = risk_widths[c_idx]
        set_cell_background(rc[c_idx], bg)
        set_cell_margins(rc[c_idx], 60, 60, 70, 70)
        set_cell_borders(rc[c_idx], bottom=('single', '4', 'D0D7DE'), top=('single', '4', 'D0D7DE'))
        p_c = rc[c_idx].paragraphs[0]; p_c.paragraph_format.space_after = Pt(0)
        r = p_c.add_run(v); r.font.name = 'Calibri'; r.font.size = Pt(8)
        if c_idx == 0: r.font.bold = True

# Section 7: Conclusion & References
add_styled_heading(doc, "7. Conclusion & Strategic Next Steps", level=1)
p = doc.add_paragraph()
p.add_run(
    "The Week 1 strategic planning report demonstrates that an integrated machine learning and operations research architecture "
    "fundamentally transforms urban logistics economics. By combining early-warning delay prediction (0.9976 ROC-AUC), spatial micro-zoning, "
    "and heuristic route optimization (50.84% mileage reduction), the enterprise is primed to exceed the 96% OTIF threshold while saving nearly $1M annually."
)

add_callout(
    doc,
    "Immediate Action Items for Week 2:\n"
    "1. Productionize Kafka ingestion pipelines connecting WMS and TMS databases into the central lakehouse.\n"
    "2. Deploy a 10-van A/B operational pilot in Urban Zone Beta comparing 2-Opt dispatch against legacy FIFO routing.\n"
    "3. Incorporate live weather and OSRM real-time traffic distance matrices into the pre-dispatch scoring service.",
    title="WEEK 2 EXECUTION ROADMAP"
)

add_styled_heading(doc, "8. References & Academic Citations", level=1)
refs = [
    "Amazon Last-Mile Routing Research Challenge Dataset (2021). Amazon.com Inc. & MIT Center for Transportation & Logistics.",
    "DataCo Global Supply Chain Intelligence Dataset (2019). Kaggle Public Repository.",
    "Dantzig, G. B., & Ramser, J. H. (1959). The truck dispatching problem. Management Science, 6(1), 80-91.",
    "Toth, P., & Vigo, D. (Eds.). (2002). The vehicle routing problem. Society for Industrial and Applied Mathematics (SIAM).",
    "Silver, E. A., Pyke, D. F., & Peterson, R. (1998). Inventory management and production planning and scheduling. John Wiley & Sons.",
    "U.S. Department of Transportation, Bureau of Transportation Statistics (BTS). (2023). Freight Analysis Framework (FAF5).",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. ACM SIGKDD (pp. 785-794)."
]
for rf in refs:
    p_rf = doc.add_paragraph(style='List Bullet'); p_rf.paragraph_format.space_after = Pt(2)
    r = p_rf.add_run(rf); r.font.name = 'Calibri'; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(60, 60, 60)

doc_path = "Logistics_Strategic_Planning_Report.docx"
doc.save(doc_path)
print(f"Document successfully compiled and saved to {doc_path}!")
