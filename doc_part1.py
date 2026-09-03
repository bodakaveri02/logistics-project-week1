import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from build_doc_report_helpers import (
    set_cell_background, set_cell_margins, set_cell_borders,
    add_styled_heading, add_callout, add_code_block, add_image_with_caption
)

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

# Cover Header
p_t = doc.add_paragraph()
p_t.paragraph_format.space_before = Pt(0)
p_t.paragraph_format.space_after = Pt(4)
r_t = p_t.add_run("STRATEGIC PLANNING & DATA EXPLORATION REPORT")
r_t.font.name = 'Arial'
r_t.font.size = Pt(21)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(27, 54, 93)

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(14)
r_sub = p_sub.add_run("Optimizing Multi-Echelon Inventory Allocation and Last-Mile Route Efficiency Through Data Science & Machine Learning")
r_sub.font.name = 'Arial'
r_sub.font.size = Pt(12)
r_sub.font.bold = True
r_sub.font.color.rgb = RGBColor(43, 92, 143)

# Metadata Table
meta_table = doc.add_table(rows=2, cols=4)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_table.autofit = False
col_widths = [Inches(1.5), Inches(1.75), Inches(1.5), Inches(1.75)]
meta_data = [
    [("Project Phase:", "Week 1: Strategy & Exploration"), ("Document Version:", "1.0 (Enterprise Plan)")],
    [("Target Network:", "Urban Omnichannel Supply Chain"), ("Core Methodology:", "ML, Spatial Clustering & VRP")]
]
for r_idx, row in enumerate(meta_data):
    c_idx = 0
    for lbl, val in row:
        cl = meta_table.cell(r_idx, c_idx)
        cv = meta_table.cell(r_idx, c_idx + 1)
        cl.width, cv.width = col_widths[c_idx], col_widths[c_idx + 1]
        set_cell_background(cl, "F0F4F8")
        set_cell_background(cv, "FFFFFF")
        set_cell_margins(cl, 60, 60, 80, 80)
        set_cell_margins(cv, 60, 60, 80, 80)
        set_cell_borders(cl, bottom=('single', '4', 'D0D7DE'), top=('single', '4', 'D0D7DE'))
        set_cell_borders(cv, bottom=('single', '4', 'D0D7DE'), top=('single', '4', 'D0D7DE'))
        pl = cl.paragraphs[0]; pl.paragraph_format.space_after = Pt(0)
        rl = pl.add_run(lbl); rl.font.name = 'Arial'; rl.font.size = Pt(8.5); rl.font.bold = True; rl.font.color.rgb = RGBColor(27, 54, 93)
        pv = cv.paragraphs[0]; pv.paragraph_format.space_after = Pt(0)
        rv = pv.add_run(val); rv.font.name = 'Calibri'; rv.font.size = Pt(8.5); rv.font.color.rgb = RGBColor(50, 50, 50)
        c_idx += 2

p_sp = doc.add_paragraph(); p_sp.paragraph_format.space_before = Pt(6); p_sp.paragraph_format.space_after = Pt(8)
