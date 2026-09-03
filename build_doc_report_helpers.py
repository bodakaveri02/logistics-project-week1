import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    """Sets internal padding (margins) for a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Sets custom borders for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style:
            val, sz, color = border_style
            b_element = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(b_element)
        else:
            b_element = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="none"/>')
            tcBorders.append(b_element)
    tcPr.append(tcBorders)

def add_styled_heading(doc, text, level):
    """Adds a heading with professional corporate styling."""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(14 if level==1 else 10)
    heading.paragraph_format.space_after = Pt(4 if level==1 else 3)
    heading.paragraph_format.keep_with_next = True
    
    # Custom color scheme
    run = heading.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(27, 54, 93)     # Deep Navy
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(43, 92, 143)    # Slate Blue
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(60, 70, 85)     # Steel Charcoal
    return heading

def add_callout(doc, text, title="STRATEGIC DIRECTIVE"):
    """Adds an executive callout box with a shaded background and navy border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=160)
    set_cell_borders(cell, left=('single', '24', '1B365D'), top=None, bottom=None, right=None)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run_t = p.add_run(f"[{title}] ")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(27, 54, 93)
    
    run_body = p.add_run(text)
    run_body.font.name = 'Calibri'
    run_body.font.size = Pt(10.5)
    run_body.font.color.rgb = RGBColor(40, 40, 40)
    
    # Empty trailing spacing
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_str, caption=""):
    """Adds a nicely formatted code block with monospace font and light gray fill."""
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after = Pt(2)
        r_cap = p_cap.add_run(f"Code Listing: {caption}")
        r_cap.font.name = 'Arial'
        r_cap.font.size = Pt(9.5)
        r_cap.font.bold = True
        r_cap.font.color.rgb = RGBColor(43, 92, 143)
        
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F7F9FB")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=140)
    set_cell_borders(cell, 
                     top=('single', '4', 'D0D7DE'), 
                     bottom=('single', '4', 'D0D7DE'), 
                     left=('single', '18', '2B5C8F'), 
                     right=('single', '4', 'D0D7DE'))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_str.strip())
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(6)

def add_image_with_caption(doc, img_path, caption, width=Inches(6.2)):
    """Inserts a centered image with a styled caption."""
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=width)
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(10)
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = 'Arial'
        run_cap.font.size = Pt(9.5)
        run_cap.font.italic = True
        run_cap.font.color.rgb = RGBColor(80, 80, 80)
    else:
        p_err = doc.add_paragraph(f"[Image file not found: {img_path}]")
        p_err.runs[0].font.color.rgb = RGBColor(200, 0, 0)

print("Helper functions defined.")
